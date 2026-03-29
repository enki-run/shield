import csv
import shutil
from docx import Document as DocxDocument
from openpyxl import Workbook
from app.pipeline.parsers.base import ParsedContent


def rebuild_document(content: ParsedContent, output_path: str, output_format: str):
    builders = {
        "txt": _rebuild_txt,
        "csv": _rebuild_csv,
        "md": _rebuild_markdown,
        "docx": _rebuild_docx,
        "xlsx": _rebuild_xlsx,
        "odt": _rebuild_odt,
        "ods": _rebuild_ods,
    }
    builder = builders.get(output_format)
    if builder is None:
        raise ValueError(f"Unsupported output format: {output_format}")
    builder(content, output_path)


def _rebuild_txt(content, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        for block in content.blocks:
            f.write(block.text + "\n")


def _rebuild_markdown(content, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        for block in content.blocks:
            if block.block_type == "table":
                f.write("\n" + block.text + "\n\n")
            else:
                f.write(block.text + "\n")


def _rebuild_csv(content, output_path):
    rows = {}
    for block in content.blocks:
        row, col = block.metadata.get("row", 0), block.metadata.get("col", 0)
        rows.setdefault(row, {})[col] = block.text
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row_idx in sorted(rows):
            cols = rows[row_idx]
            max_col = max(cols) if cols else 0
            writer.writerow([cols.get(c, "") for c in range(max_col + 1)])


def _rebuild_docx(content, output_path):
    """In-place replacement: copy original, replace PII strings in runs."""
    original_path = content.metadata.get("original_path")
    replacements = content.metadata.get("replacements", {})
    if not original_path or not replacements:
        # Fallback: create new doc from blocks
        doc = DocxDocument()
        for block in content.blocks:
            doc.add_paragraph(block.text)
        doc.save(output_path)
        return
    shutil.copy2(original_path, output_path)
    doc = DocxDocument(output_path)
    for para in doc.paragraphs:
        for run in para.runs:
            for original, pseudonym in replacements.items():
                if original in run.text:
                    run.text = run.text.replace(original, pseudonym)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for original, pseudonym in replacements.items():
                            if original in run.text:
                                run.text = run.text.replace(original, pseudonym)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for para in hf.paragraphs:
                    for run in para.runs:
                        for original, pseudonym in replacements.items():
                            if original in run.text:
                                run.text = run.text.replace(original, pseudonym)
    doc.save(output_path)


def _rebuild_xlsx(content, output_path):
    wb = Workbook()
    sheets = {}
    for block in content.blocks:
        sheet_name = block.metadata.get("sheet", "Sheet1")
        if sheet_name not in sheets:
            ws = wb.active if not sheets else wb.create_sheet(sheet_name)
            if not sheets:
                ws.title = sheet_name
            sheets[sheet_name] = ws
        ws = sheets[sheet_name]
        ws.cell(
            row=block.metadata.get("row", 0) + 1,
            column=block.metadata.get("col", 0) + 1,
            value=block.text,
        )
    wb.save(output_path)


def _rebuild_odt(content, output_path):
    from odf.opendocument import OpenDocumentText
    from odf.text import P

    doc = OpenDocumentText()
    for block in content.blocks:
        doc.text.addElement(P(text=block.text))
    doc.save(output_path)


def _rebuild_ods(content, output_path):
    from odf.opendocument import OpenDocumentSpreadsheet
    from odf.table import Table, TableRow, TableCell
    from odf.text import P

    doc = OpenDocumentSpreadsheet()
    rows_data = {}
    for block in content.blocks:
        sheet = block.metadata.get("sheet", "Sheet1")
        row, col = block.metadata.get("row", 0), block.metadata.get("col", 0)
        rows_data.setdefault(sheet, {}).setdefault(row, {})[col] = block.text
    for sheet_name, rows in rows_data.items():
        table = Table(name=sheet_name)
        for row_idx in sorted(rows):
            tr = TableRow()
            cols = rows[row_idx]
            for col_idx in range(max(cols) + 1 if cols else 0):
                tc = TableCell()
                text = cols.get(col_idx, "")
                if text:
                    tc.addElement(P(text=text))
                tr.addElement(tc)
            table.addElement(tr)
        doc.spreadsheet.addElement(table)
    doc.save(output_path)
