from docx import Document

from .base import BaseParser, ContentBlock, ParsedContent


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> ParsedContent:
        doc = Document(file_path)
        blocks = []

        for para in doc.paragraphs:
            if para.text.strip():
                blocks.append(ContentBlock(text=para.text, block_type="paragraph"))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        blocks.append(
                            ContentBlock(text=cell.text, block_type="table_cell")
                        )

        return ParsedContent(blocks=blocks, format="docx")
