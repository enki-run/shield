import csv
import os
import pytest
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.pipeline.parsers.base import ContentBlock, ParsedContent
from app.pipeline.rebuilder import rebuild_document


def make_content(blocks, metadata=None):
    return ParsedContent(blocks=blocks, metadata=metadata or {})


# ---------------------------------------------------------------------------
# 1. TXT
# ---------------------------------------------------------------------------

def test_rebuild_txt(tmp_path):
    blocks = [
        ContentBlock(text="Hallo Welt"),
        ContentBlock(text="Zweite Zeile"),
    ]
    content = make_content(blocks)
    out = str(tmp_path / "output.txt")
    rebuild_document(content, out, "txt")

    with open(out, encoding="utf-8") as f:
        lines = f.read().splitlines()

    assert lines[0] == "Hallo Welt"
    assert lines[1] == "Zweite Zeile"


# ---------------------------------------------------------------------------
# 2. CSV
# ---------------------------------------------------------------------------

def test_rebuild_csv(tmp_path):
    blocks = [
        ContentBlock(text="Name",   metadata={"row": 0, "col": 0}),
        ContentBlock(text="Email",  metadata={"row": 0, "col": 1}),
        ContentBlock(text="PERSON-A7F3", metadata={"row": 1, "col": 0}),
        ContentBlock(text="EMAIL-38B2",  metadata={"row": 1, "col": 1}),
    ]
    content = make_content(blocks)
    out = str(tmp_path / "output.csv")
    rebuild_document(content, out, "csv")

    with open(out, newline="", encoding="utf-8") as f:
        rows = list(csv.reader(f))

    assert rows[0] == ["Name", "Email"]
    assert rows[1] == ["PERSON-A7F3", "EMAIL-38B2"]


# ---------------------------------------------------------------------------
# 3. Markdown (pdf → md rebuild)
# ---------------------------------------------------------------------------

def test_rebuild_pdf_as_markdown(tmp_path):
    blocks = [
        ContentBlock(text="# Überschrift"),
        ContentBlock(text="Ein normaler Absatz."),
        ContentBlock(text="| A | B |\n|---|---|\n| 1 | 2 |", block_type="table"),
    ]
    content = make_content(blocks)
    out = str(tmp_path / "output.md")
    rebuild_document(content, out, "md")

    with open(out, encoding="utf-8") as f:
        text = f.read()

    assert "# Überschrift" in text
    assert "Ein normaler Absatz." in text
    assert "| A | B |" in text


# ---------------------------------------------------------------------------
# 4. DOCX in-place replacement
# ---------------------------------------------------------------------------

def test_rebuild_docx_in_place(tmp_path):
    # Create a DOCX with PII content
    original = str(tmp_path / "original.docx")
    doc = DocxDocument()
    doc.add_paragraph("Kontakt: Max Müller, max@example.com")
    doc.add_paragraph("Weitere Infos folgen.")
    doc.save(original)

    replacements = {
        "Max Müller": "PERSON-A7F3",
        "max@example.com": "EMAIL-38B2",
    }
    content = make_content(
        blocks=[],
        metadata={"original_path": original, "replacements": replacements},
    )

    out = str(tmp_path / "output.docx")
    rebuild_document(content, out, "docx")

    result_doc = DocxDocument(out)
    full_text = "\n".join(p.text for p in result_doc.paragraphs)

    assert "PERSON-A7F3" in full_text
    assert "EMAIL-38B2" in full_text
    assert "Max Müller" not in full_text
    assert "max@example.com" not in full_text


# ---------------------------------------------------------------------------
# 5. XLSX
# ---------------------------------------------------------------------------

def test_rebuild_xlsx(tmp_path):
    blocks = [
        ContentBlock(text="Vorname",     metadata={"sheet": "Sheet1", "row": 0, "col": 0}),
        ContentBlock(text="Nachname",    metadata={"sheet": "Sheet1", "row": 0, "col": 1}),
        ContentBlock(text="PERSON-A7F3", metadata={"sheet": "Sheet1", "row": 1, "col": 0}),
        ContentBlock(text="PERSON-B2C1", metadata={"sheet": "Sheet1", "row": 1, "col": 1}),
    ]
    content = make_content(blocks)
    out = str(tmp_path / "output.xlsx")
    rebuild_document(content, out, "xlsx")

    wb = load_workbook(out)
    ws = wb.active

    assert ws.cell(row=1, column=1).value == "Vorname"
    assert ws.cell(row=1, column=2).value == "Nachname"
    assert ws.cell(row=2, column=1).value == "PERSON-A7F3"
    assert ws.cell(row=2, column=2).value == "PERSON-B2C1"


# ---------------------------------------------------------------------------
# 6. Unsupported format raises ValueError
# ---------------------------------------------------------------------------

def test_rebuild_unsupported_format(tmp_path):
    content = make_content([ContentBlock(text="test")])
    with pytest.raises(ValueError, match="Unsupported output format"):
        rebuild_document(content, str(tmp_path / "out.xyz"), "xyz")
